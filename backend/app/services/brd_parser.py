import re
from typing import Dict, List, Any, Optional
from docx import Document


# Required section categories for parse_quality reporting.
# Fuzzy-matched: any section title containing these keywords qualifies.
_REQUIRED_SECTION_KEYWORDS = [
    "project overview",
    "business objective",
    "functional requirement",
    "non-functional requirement",
]


def _serialize_table_to_text(headers: List[str], rows: List[Dict]) -> str:
    """Render a structured table as markdown text for inclusion in raw_text."""
    if not headers:
        return ""
    lines = ["| " + " | ".join(headers) + " |"]
    lines.append("|" + "|".join(["---"] * len(headers)) + "|")
    for row in rows:
        lines.append("| " + " | ".join(str(row.get(h, "")) for h in headers) + " |")
    return "\n".join(lines)


class BRDParser:
    """Parse BRD documents (.docx and .txt) into raw_text, sections, tables, and parse_quality."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        if file_path.endswith('.docx'):
            return self._parse_docx(file_path)
        elif file_path.endswith('.txt'):
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    # ─────────────────────────────────────────────────────────────
    # .docx
    # ─────────────────────────────────────────────────────────────

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)

        # Build element → object maps so we can iterate body children in document order.
        # This ensures tables appear in raw_text at the correct position (not after all paragraphs).
        para_map = {p._p: p for p in doc.paragraphs}
        table_map = {t._tbl: t for t in doc.tables}

        raw_text_parts: List[str] = []
        sections: List[Dict] = []
        tables: List[Dict] = []
        current_section: Optional[Dict] = None
        table_counter = 0

        for child in doc.element.body:
            if child in para_map:
                para = para_map[child]
                text = para.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)
                heading_level = self._docx_heading_level(para)

                if heading_level:
                    if current_section:
                        sections.append(current_section)
                    current_section = {"title": text, "text": ""}
                else:
                    if current_section:
                        current_section["text"] += text + "\n"

            elif child in table_map:
                table = table_map[child]
                table_counter += 1
                parsed = self._parse_docx_table(table, table_counter)
                if parsed:
                    table_text = _serialize_table_to_text(parsed["headers"], parsed["rows"])
                    raw_text_parts.append(table_text)
                    tables.append(parsed)
                    if current_section:
                        current_section["text"] += table_text + "\n"

        if current_section:
            sections.append(current_section)

        sections = [s for s in sections if s["text"].strip()]
        raw_text = "\n\n".join(raw_text_parts)

        return {
            "raw_text": raw_text,
            "sections": sections,
            "tables": tables,
            "parse_quality": self._build_quality(raw_text, sections, tables),
        }

    def _docx_heading_level(self, para) -> Optional[int]:
        """Return heading level (1–6) or None."""
        style_name = para.style.name if para.style else ""

        # Primary: Word heading styles
        if style_name.lower().startswith('heading'):
            parts = style_name.split()
            try:
                return int(parts[-1])
            except (ValueError, IndexError):
                return 1

        # Fallback: short paragraph where every run with text is bold
        text = para.text.strip()
        if not text or len(text) > 120 or text[-1] in '.?!':
            return None

        runs_with_text = [r for r in para.runs if r.text.strip()]
        if runs_with_text and all(r.bold for r in runs_with_text):
            return 2  # treat bold-only paragraphs as H2

        return None

    def _parse_docx_table(self, table, counter: int) -> Optional[Dict]:
        """Extract headers and rows from a python-docx Table object."""
        if len(table.rows) < 2:
            return None

        # Extract headers from row 0, deduplicating merged cells
        raw_headers = []
        prev_cell = None
        for cell in table.rows[0].cells:
            if cell is prev_cell:
                raw_headers.append("")  # merged — keep slot empty
            else:
                raw_headers.append(cell.text.strip())
            prev_cell = cell

        if not any(raw_headers):
            return None

        rows = []
        for row in table.rows[1:]:
            row_data = {}
            prev_cell = None
            for idx, cell in enumerate(row.cells):
                if cell is prev_cell:
                    idx += 1
                    continue
                prev_cell = cell
                if idx < len(raw_headers) and raw_headers[idx]:
                    row_data[raw_headers[idx]] = cell.text.strip()
            if row_data:
                rows.append(row_data)

        headers = [h for h in raw_headers if h]
        if not headers:
            return None

        return {
            "table_id": f"table_{counter}",
            "headers": headers,
            "rows": rows,
        }

    # ─────────────────────────────────────────────────────────────
    # .txt
    # ─────────────────────────────────────────────────────────────

    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        content = self._read_with_fallback(file_path)
        lines = content.splitlines()

        sections: List[Dict] = []
        tables: List[Dict] = []
        current_section: Optional[Dict] = None
        table_counter = 0
        i = 0

        while i < len(lines):
            stripped = lines[i].strip()

            if not stripped:
                i += 1
                continue

            # Detect markdown table: header row followed by separator row (---|---|)
            if (
                '|' in stripped
                and i + 1 < len(lines)
                and re.match(r'^[\|\s\-:]+$', lines[i + 1].strip())
            ):
                table_lines = []
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                parsed = self._parse_markdown_table(table_lines, table_counter + 1)
                if parsed:
                    table_counter += 1
                    tables.append(parsed)
                    table_text = _serialize_table_to_text(parsed["headers"], parsed["rows"])
                    if current_section:
                        current_section["text"] += table_text + "\n"
                continue

            # Detect heading
            heading_level = self._txt_heading_level(stripped)
            if heading_level:
                if current_section:
                    sections.append(current_section)
                current_section = {"title": stripped, "text": ""}
            else:
                if current_section:
                    current_section["text"] += stripped + "\n"

            i += 1

        if current_section:
            sections.append(current_section)

        sections = [s for s in sections if s["text"].strip()]

        return {
            "raw_text": content,
            "sections": sections,
            "tables": tables,
            "parse_quality": self._build_quality(content, sections, tables),
        }

    def _read_with_fallback(self, file_path: str) -> str:
        """Read text file trying utf-8, utf-8-sig, then latin-1."""
        for encoding in ('utf-8', 'utf-8-sig', 'latin-1'):
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode {file_path}: unsupported encoding")

    def _txt_heading_level(self, line: str) -> Optional[int]:
        """
        Detect headings in .txt files.
        Primary signal: numbered sections (1., 1.1., 1.1.1.)
        Fallback: markdown headings (#, ##, ###)
        """
        # Numbered: 1. Title, 1.1 Title, 1.1.1 Title
        m = re.match(r'^(\d+(?:\.\d+)*\.?)\s+\S', line)
        if m:
            depth = m.group(1).rstrip('.').count('.') + 1
            return min(depth, 3)

        # Markdown: # H1, ## H2, ### H3
        m = re.match(r'^(#{1,3})\s+\S', line)
        if m:
            return len(m.group(1))

        return None

    def _parse_markdown_table(self, lines: List[str], counter: int) -> Optional[Dict]:
        """Parse markdown table lines into structured data."""
        if len(lines) < 3:  # need header + separator + at least one row
            return None

        def split_row(line: str) -> List[str]:
            return [cell.strip() for cell in line.strip().strip('|').split('|')]

        headers = split_row(lines[0])
        if not any(headers):
            return None

        rows = []
        for line in lines[2:]:  # skip header + separator
            cells = split_row(line)
            row = {h: cells[i] for i, h in enumerate(headers) if h and i < len(cells)}
            if row:
                rows.append(row)

        return {
            "table_id": f"table_{counter}",
            "headers": [h for h in headers if h],
            "rows": rows,
        }

    # ─────────────────────────────────────────────────────────────
    # Parse quality
    # ─────────────────────────────────────────────────────────────

    def _build_quality(
        self, raw_text: str, sections: List[Dict], tables: List[Dict]
    ) -> Dict:
        titles_lower = [s["title"].lower() for s in sections]
        missing = [
            kw for kw in _REQUIRED_SECTION_KEYWORDS
            if not any(kw in t for t in titles_lower)
        ]
        return {
            "word_count": len(raw_text.split()),
            "char_count": len(raw_text),
            "section_count": len(sections),
            "table_count": len(tables),
            "missing_required_sections": missing,
        }
